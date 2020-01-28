import azure.cognitiveservices.speech as speechsdk
import tkinter as tk
import time
import docx
from tkinter import *
from multiprocessing import Process
import threading




speech_key, service_region = "YOUR_KEY", "YOUR_REGION"
speech_config = speechsdk.SpeechConfig(subscription=speech_key, region=service_region)
speech_recognizer = speechsdk.SpeechRecognizer(speech_config=speech_config)


doc = docx.Document()
header = input('Enter your title :')
# add a heading of level 0 (largest heading)
doc.add_heading(header, 0)
# add a paragraph and store
# the object in a variable
doc_para = doc.add_paragraph(' ')
# add a run i.e, style like
# bold, italic, underline, etc.
doc_para.add_run('hey there, bold here').bold = True
#doc_para.add_run(', and ')
#doc_para.add_run('these words are italic').italic = True
space = ' '
#word = ['you' + space, 'are' + space, str(1) + space]

# Connect callbacks to the events fired by the speech recognizer
# a=""
class A:
    def __init__(self, master):
        self.label=tk.Label(master)
        self.label.grid(row=0, column=0)
        self.label.configure(text='nothing')
        self.text = ''
        self.count1 = 0
        self.update_label()

    def swap(self,a):
        temp = self.text
        self.text = a
        a = temp
        return a

    def update_label(self):
        word = self.swap(a)
        self.label.configure(text = word)
        self.label.after(1000, self.update_label) # call this method again in 1,000 milliseconds
        self.count1 += 1

'''window=tk.Tk()
window'''


def start():
    root = tk.Tk()
    A(root).update_label()
    time.sleep(2)

    root.mainloop()

count = 0
while count <= 1:
    result=speech_recognizer.recognize_once()
    speech_recognizer.recognized.connect(print('RECOGNIZED: {}'.format(result.text)))
    a=result.text
    start()


   
 # stop continuous recognition on either session stopped or canceled events


   
doc.save('text.doc')

speech_recognizer.stop_continuous_recognition()
